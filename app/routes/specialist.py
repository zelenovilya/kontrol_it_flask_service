from pathlib import Path
from datetime import datetime
from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app, send_file
from flask_login import login_required, current_user
from docx import Document as DocxDocument
from openpyxl import Workbook
from app import db
from app.decorators import role_required
from app.models import Ticket, TicketStatus, TicketComment, Attachment, Document, Employee
from app.utils import save_uploaded_file

specialist_bp = Blueprint("specialist", __name__, url_prefix="/specialist")


def specialist_breadcrumbs(*items):
    trail = [{"label": "Главная", "url": url_for("public.home")}]
    if items:
        trail.append({"label": "Кабинет специалиста", "url": url_for("specialist.dashboard")})
        trail.extend(items)
    else:
        trail.append({"label": "Кабинет специалиста", "url": None})
    return trail


@specialist_bp.route("/dashboard")
@login_required
@role_required("Специалист")
def dashboard():
    employee = current_user.employee
    tickets = Ticket.query.filter_by(specialist=employee).order_by(Ticket.updated_at.desc()).all()
    return render_template(
        "specialist/dashboard.html",
        title="Панель специалиста",
        tickets=tickets,
        breadcrumbs=specialist_breadcrumbs(),
    )


@specialist_bp.route("/new-tickets")
@login_required
@role_required("Специалист")
def new_tickets():
    new_status = TicketStatus.query.filter_by(name="Новая").first()
    tickets = Ticket.query.filter_by(status=new_status).order_by(Ticket.created_at.asc()).all()
    return render_template(
        "specialist/new_tickets.html",
        title="Новые заявки",
        tickets=tickets,
        breadcrumbs=specialist_breadcrumbs({"label": "Новые заявки", "url": None}),
    )


@specialist_bp.route("/tickets/<int:ticket_id>", methods=["GET", "POST"])
@login_required
@role_required("Специалист")
def ticket_detail(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    statuses = TicketStatus.query.order_by(TicketStatus.sort_order).all()
    specialists = Employee.query.all()
    if request.method == "POST":
        ticket.status_id = int(request.form.get("status_id"))
        specialist_id = request.form.get("specialist_id")
        ticket.specialist_id = int(specialist_id) if specialist_id else current_user.employee.id
        comment = request.form.get("comment", "").strip()
        if comment:
            db.session.add(TicketComment(ticket=ticket, author_id=current_user.id, text=comment))
        uploaded = save_uploaded_file(request.files.get("file"), current_app.config["UPLOAD_FOLDER"], prefix=f"ticket_{ticket.id}")
        if uploaded:
            original_name, stored_name, file_path = uploaded
            db.session.add(Attachment(original_name=original_name, stored_name=stored_name, file_path=file_path, ticket=ticket, uploaded_by_id=current_user.id))
        db.session.commit()
        flash("Карточка заявки обновлена.", "success")
        return redirect(url_for("specialist.ticket_detail", ticket_id=ticket.id))
    return render_template(
        "specialist/ticket_detail.html",
        title=f"Заявка №{ticket.id}",
        ticket=ticket,
        statuses=statuses,
        specialists=specialists,
        breadcrumbs=specialist_breadcrumbs({"label": f"Заявка №{ticket.id}", "url": None}),
    )


@specialist_bp.route("/tickets/<int:ticket_id>/history")
@login_required
@role_required("Специалист")
def ticket_history(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    return render_template(
        "specialist/history.html",
        title=f"История заявки №{ticket.id}",
        ticket=ticket,
        breadcrumbs=specialist_breadcrumbs(
            {"label": f"Заявка №{ticket.id}", "url": url_for("specialist.ticket_detail", ticket_id=ticket.id)},
            {"label": "История работ", "url": None},
        ),
    )


@specialist_bp.route("/tickets/<int:ticket_id>/act")
@login_required
@role_required("Специалист")
def act_page(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    return render_template(
        "specialist/act.html",
        title=f"Акт по заявке №{ticket.id}",
        ticket=ticket,
        breadcrumbs=specialist_breadcrumbs(
            {"label": f"Заявка №{ticket.id}", "url": url_for("specialist.ticket_detail", ticket_id=ticket.id)},
            {"label": "Формирование акта", "url": None},
        ),
    )


@specialist_bp.route("/tickets/<int:ticket_id>/generate-act")
@login_required
@role_required("Специалист")
def generate_act(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    doc = DocxDocument()
    doc.add_heading("Акт выполненных работ", level=1)
    doc.add_paragraph(f"Номер заявки: {ticket.id}")
    doc.add_paragraph(f"Клиент: {ticket.client.organization_name}")
    doc.add_paragraph(f"Услуга: {ticket.service.name}")
    doc.add_paragraph(f"Статус: {ticket.status.name}")
    doc.add_paragraph(f"Описание заявки: {ticket.description}")
    doc.add_paragraph(f"Исполнитель: {ticket.specialist.user.full_name if ticket.specialist else current_user.full_name}")
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph("Работы выполнены в соответствии с заявкой клиента. Претензии по объему выполненных работ отсутствуют.")

    file_name = f"act_ticket_{ticket.id}.docx"
    file_path = Path(current_app.config["GENERATED_DOCS_FOLDER"]) / file_name
    doc.save(file_path)

    document = Document(title=f"Акт по заявке №{ticket.id}", document_type="Акт", file_path=str(file_path), ticket=ticket, created_by_id=current_user.id)
    db.session.add(document)
    db.session.commit()
    flash("Акт выполненных работ сформирован.", "success")
    return send_file(file_path, as_attachment=True)


@specialist_bp.route("/reports")
@login_required
@role_required("Специалист")
def reports():
    completed_statuses = TicketStatus.query.filter(TicketStatus.name.in_(["Выполнена", "Закрыта"])).all()
    completed_status_ids = [status.id for status in completed_statuses]
    tickets = Ticket.query.filter(
        Ticket.specialist == current_user.employee,
        Ticket.status_id.in_(completed_status_ids),
    ).order_by(Ticket.updated_at.desc()).all()
    return render_template(
        "specialist/reports.html",
        title="Отчет по выполненным заявкам",
        tickets=tickets,
        breadcrumbs=specialist_breadcrumbs({"label": "Отчетность", "url": None}),
    )


@specialist_bp.route("/documents")
@login_required
@role_required("Специалист")
def documents():
    docs = (
        Document.query.join(Ticket)
        .filter(Ticket.specialist_id == current_user.employee.id)
        .order_by(Document.created_at.desc())
        .all()
    )
    return render_template(
        "specialist/documents.html",
        title="Документы специалиста",
        documents=docs,
        breadcrumbs=specialist_breadcrumbs({"label": "Документы", "url": None}),
    )


@specialist_bp.route("/reports/tickets.xlsx")
@login_required
@role_required("Специалист")
def tickets_report():
    wb = Workbook()
    ws = wb.active
    ws.title = "Заявки"
    ws.append(["№", "Дата", "Клиент", "Услуга", "Статус", "Приоритет", "Исполнитель"])
    for ticket in Ticket.query.filter_by(specialist=current_user.employee).order_by(Ticket.created_at.desc()).all():
        ws.append([
            ticket.id,
            ticket.created_at.strftime("%d.%m.%Y %H:%M"),
            ticket.client.organization_name,
            ticket.service.name,
            ticket.status.name,
            ticket.priority.name,
            ticket.specialist.user.full_name if ticket.specialist else "Не назначен",
        ])
    file_path = Path(current_app.config["GENERATED_REPORTS_FOLDER"]) / "tickets_report.xlsx"
    wb.save(file_path)
    return send_file(file_path, as_attachment=True)


@specialist_bp.route("/help")
@login_required
@role_required("Специалист")
def help_page():
    return render_template(
        "specialist/help.html",
        title="Справка специалиста",
        breadcrumbs=specialist_breadcrumbs({"label": "Справка специалиста", "url": None}),
    )
