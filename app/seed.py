from datetime import datetime, timedelta
from pathlib import Path

from docx import Document as DocxDocument
from flask import current_app
from werkzeug.security import generate_password_hash

from app import db
from app.models import (
    Article,
    Attachment,
    Client,
    Document,
    Employee,
    FeedbackMessage,
    Role,
    Service,
    ServiceCategory,
    Ticket,
    TicketComment,
    TicketPriority,
    TicketStatus,
    User,
)


def write_sample_file(path, title, lines):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(title + "\n" + "\n".join(lines) + "\n", encoding="utf-8")


def write_sample_act(path, ticket_number, client_name, service_name, specialist_name):
    path.parent.mkdir(parents=True, exist_ok=True)
    document = DocxDocument()
    document.add_heading("Акт выполненных работ", level=1)
    document.add_paragraph(f"Номер заявки: {ticket_number}")
    document.add_paragraph(f"Клиент: {client_name}")
    document.add_paragraph(f"Услуга: {service_name}")
    document.add_paragraph(f"Исполнитель: {specialist_name}")
    document.add_paragraph("Работы выполнены в рамках учебного веб-сервиса ООО «Контроль ИТ».")
    document.add_paragraph("Документ создан как тестовые данные для проверки скачивания DOCX.")
    document.save(path)


def init_db():
    db.create_all()

    if Role.query.count() > 0:
        return

    admin_role = Role(name="Администратор", description="Полный доступ к управлению сервисом")
    specialist_role = Role(name="Специалист", description="Обработка и выполнение заявок")
    client_role = Role(name="Клиент", description="Создание заявок и просмотр своих обращений")
    db.session.add_all([admin_role, specialist_role, client_role])
    db.session.flush()

    admin = User(
        username="admin",
        password_hash=generate_password_hash("admin123", method="pbkdf2:sha256"),
        email="admin@kontrol-it.local",
        full_name="Администратор системы",
        phone="+7 900 000-00-01",
        role=admin_role,
    )
    specialist_user = User(
        username="specialist",
        password_hash=generate_password_hash("spec123", method="pbkdf2:sha256"),
        email="specialist@kontrol-it.local",
        full_name="Иванов Сергей Петрович",
        phone="+7 900 000-00-02",
        role=specialist_role,
    )
    specialist_user_2 = User(
        username="engineer",
        password_hash=generate_password_hash("engineer123", method="pbkdf2:sha256"),
        email="engineer@kontrol-it.local",
        full_name="Смирнова Анна Викторовна",
        phone="+7 900 000-00-04",
        role=specialist_role,
    )
    client_user = User(
        username="client",
        password_hash=generate_password_hash("client123", method="pbkdf2:sha256"),
        email="client@example.local",
        full_name="Петров Алексей Николаевич",
        phone="+7 900 000-00-03",
        role=client_role,
    )
    client_user_2 = User(
        username="client2",
        password_hash=generate_password_hash("client2123", method="pbkdf2:sha256"),
        email="office@example.local",
        full_name="Кузнецова Мария Андреевна",
        phone="+7 900 000-00-05",
        role=client_role,
    )
    db.session.add_all([admin, specialist_user, specialist_user_2, client_user, client_user_2])
    db.session.flush()

    employee = Employee(
        user=specialist_user,
        position="Ведущий специалист технической поддержки",
        department="Отдел сопровождения клиентов",
        qualification="Сетевое администрирование, сопровождение рабочих мест",
    )
    employee_2 = Employee(
        user=specialist_user_2,
        position="Инженер по инфраструктуре",
        department="Отдел серверных решений",
        qualification="Windows Server, резервное копирование, учетные записи",
    )
    client = Client(
        user=client_user,
        organization_name="ООО «Демо-Клиент»",
        inn="7700000000",
        address="г. Москва, ул. Примерная, д. 1",
        contact_person="Петров Алексей Николаевич",
    )
    client_2 = Client(
        user=client_user_2,
        organization_name="АО «Офисные решения»",
        inn="7800000000",
        address="г. Санкт-Петербург, Невский проспект, д. 10",
        contact_person="Кузнецова Мария Андреевна",
    )
    db.session.add_all([employee, employee_2, client, client_2])

    categories = [
        ServiceCategory(name="Техническая поддержка", description="Оперативное решение пользовательских инцидентов"),
        ServiceCategory(name="Администрирование", description="Сопровождение серверов, рабочих станций и сетевой инфраструктуры"),
        ServiceCategory(name="Информационная безопасность", description="Проверка доступа, резервное копирование, базовая защита"),
        ServiceCategory(name="Документы и отчетность", description="Подготовка актов, выгрузок и подтверждающих документов"),
    ]
    db.session.add_all(categories)
    db.session.flush()

    services = [
        Service(name="Настройка рабочего места", description="Подключение пользователя, ПО, принтера и сетевых ресурсов", price=2500, category=categories[0]),
        Service(name="Удаленная диагностика инцидента", description="Проверка проблемы через удаленное подключение", price=1500, category=categories[0]),
        Service(name="Настройка резервного копирования", description="Подготовка расписания и проверка восстановления", price=5000, category=categories[2]),
        Service(name="Сопровождение серверов", description="Обновление, мониторинг и регламентное обслуживание", price=12000, category=categories[1]),
        Service(name="Аудит учетных записей", description="Проверка ролей, прав доступа и неактивных пользователей", price=7000, category=categories[2]),
        Service(name="Настройка VPN-доступа", description="Подключение удаленного сотрудника к корпоративным ресурсам", price=4200, category=categories[1]),
        Service(name="Подготовка акта работ", description="Формирование подтверждающего документа по выполненной заявке", price=900, category=categories[3]),
        Service(name="Проверка офисного принтера", description="Диагностика печати, драйверов и сетевого подключения", price=1800, category=categories[0]),
    ]
    db.session.add_all(services)

    statuses = [
        TicketStatus(name="Новая", sort_order=1),
        TicketStatus(name="Принята", sort_order=2),
        TicketStatus(name="В работе", sort_order=3),
        TicketStatus(name="Ожидает клиента", sort_order=4),
        TicketStatus(name="Выполнена", sort_order=5),
        TicketStatus(name="Закрыта", sort_order=6),
    ]
    priorities = [
        TicketPriority(name="Низкий", response_hours=48),
        TicketPriority(name="Средний", response_hours=24),
        TicketPriority(name="Высокий", response_hours=8),
        TicketPriority(name="Критический", response_hours=2),
    ]
    db.session.add_all(statuses + priorities)
    db.session.flush()

    now = datetime.utcnow()
    tickets = [
        Ticket(
            title="Не открывается корпоративная почта",
            description="При входе в почтовый ящик появляется сообщение об ошибке авторизации.",
            client=client,
            specialist=employee,
            service=services[1],
            status=statuses[2],
            priority=priorities[1],
            deadline_at=now + timedelta(hours=24),
        ),
        Ticket(
            title="Нужно настроить новое рабочее место",
            description="Для нового сотрудника требуется установить офисное ПО, подключить принтер и сетевые папки.",
            client=client,
            specialist=employee,
            service=services[0],
            status=statuses[1],
            priority=priorities[1],
            deadline_at=now + timedelta(hours=24),
        ),
        Ticket(
            title="Проверить резервное копирование бухгалтерской базы",
            description="Необходимо убедиться, что резервные копии создаются и восстанавливаются без ошибок.",
            client=client_2,
            specialist=employee_2,
            service=services[2],
            status=statuses[4],
            priority=priorities[2],
            deadline_at=now + timedelta(hours=8),
        ),
        Ticket(
            title="Настроить VPN для удаленного отдела",
            description="Пять сотрудников работают удаленно и должны получить защищенный доступ к файловому серверу.",
            client=client_2,
            specialist=employee_2,
            service=services[5],
            status=statuses[0],
            priority=priorities[2],
            deadline_at=now + timedelta(hours=8),
        ),
        Ticket(
            title="Периодически пропадает сетевой принтер",
            description="Печать зависает после нескольких заданий, пользователи перезапускают принтер вручную.",
            client=client,
            specialist=None,
            service=services[7],
            status=statuses[0],
            priority=priorities[0],
            deadline_at=now + timedelta(hours=48),
        ),
        Ticket(
            title="Сформировать акт по завершенным работам",
            description="Нужен акт по выполненной диагностике и настройке почтового клиента.",
            client=client,
            specialist=employee,
            service=services[6],
            status=statuses[5],
            priority=priorities[0],
            deadline_at=now - timedelta(hours=4),
            closed_at=now - timedelta(hours=1),
        ),
    ]
    db.session.add_all(tickets)
    db.session.flush()

    comments = [
        TicketComment(ticket=tickets[0], author=client_user, text="Проблема повторяется после перезагрузки компьютера."),
        TicketComment(ticket=tickets[0], author=specialist_user, text="Заявка принята, выполняется проверка учетной записи."),
        TicketComment(ticket=tickets[1], author=admin, text="Назначен специалист технической поддержки."),
        TicketComment(ticket=tickets[2], author=specialist_user_2, text="Проверка выполнена, восстановление тестовой копии прошло успешно."),
        TicketComment(ticket=tickets[3], author=client_user_2, text="Список сотрудников отправлен в приложенном файле."),
        TicketComment(ticket=tickets[5], author=specialist_user, text="Работы завершены, акт сформирован."),
    ]
    db.session.add_all(comments)

    upload_folder = Path(current_app.config["UPLOAD_FOLDER"])
    attachment_1 = upload_folder / "sample_ticket_1_mail_error.txt"
    attachment_2 = upload_folder / "sample_ticket_4_vpn_users.txt"
    write_sample_file(
        attachment_1,
        "Диагностический файл по заявке №1",
        [
            "Ошибка авторизации возникает у пользователя demo.client.",
            "Проверить пароль, срок действия учетной записи и доступ к почтовому серверу.",
        ],
    )
    write_sample_file(
        attachment_2,
        "Список сотрудников для VPN",
        [
            "Иванова Елена, отдел продаж",
            "Соколов Павел, отдел продаж",
            "Орлова Ирина, бухгалтерия",
        ],
    )
    db.session.add_all([
        Attachment(
            original_name="mail_error_diagnostics.txt",
            stored_name=attachment_1.name,
            file_path=str(attachment_1),
            ticket=tickets[0],
            uploaded_by=client_user,
        ),
        Attachment(
            original_name="vpn_users.txt",
            stored_name=attachment_2.name,
            file_path=str(attachment_2),
            ticket=tickets[3],
            uploaded_by=client_user_2,
        ),
    ])

    generated_docs_folder = Path(current_app.config["GENERATED_DOCS_FOLDER"])
    sample_act = generated_docs_folder / "sample_act_ticket_6.docx"
    write_sample_act(sample_act, tickets[5].id, client.organization_name, services[6].name, specialist_user.full_name)
    db.session.add(
        Document(
            title="Акт по заявке №6",
            document_type="Акт",
            file_path=str(sample_act),
            ticket=tickets[5],
            created_by=specialist_user,
        )
    )

    feedback_messages = [
        FeedbackMessage(
            name="Орлова Ирина",
            email="orlovai@example.local",
            phone="+7 900 100-00-10",
            subject="Вопрос по регистрации",
            message="Подскажите, можно ли зарегистрировать несколько контактных лиц от одной организации?",
            is_processed=False,
        ),
        FeedbackMessage(
            name="ООО «Партнер»",
            email="partner@example.local",
            phone="+7 900 100-00-11",
            subject="Уточнение по тарифам",
            message="Интересует порядок учета срочных заявок и формирования отчетов по месяцам.",
            is_processed=True,
        ),
    ]
    db.session.add_all(feedback_messages)

    articles = [
        Article(
            title="Как правильно создать заявку",
            slug="how-to-create-ticket",
            summary="Краткая инструкция для клиентов",
            content="Опишите проблему, выберите услугу и при необходимости прикрепите файл. Чем подробнее заявка, тем быстрее специалист сможет приступить к работе.",
            article_type="knowledge",
        ),
        Article(
            title="Регламент обработки обращений",
            slug="ticket-processing-rules",
            summary="Основные статусы и сроки",
            content="Заявки проходят статусы от новой до закрытой. Срок реакции зависит от приоритета, а все изменения фиксируются в комментариях.",
            article_type="knowledge",
        ),
        Article(
            title="Работа с файлами в заявке",
            slug="ticket-files-rules",
            summary="Кто может загружать и скачивать вложения",
            content="Клиент скачивает файлы только по своим заявкам, специалист — по назначенным, администратор — по всем обращениям.",
            article_type="knowledge",
        ),
        Article(
            title="ООО «Контроль ИТ» запускает клиентский сервис",
            slug="service-launch",
            summary="Новость о внедрении веб-сервиса",
            content="Единый сервис повышает прозрачность обработки заявок и упрощает контроль работ.",
            article_type="news",
        ),
        Article(
            title="Добавлена выгрузка отчетов",
            slug="reports-export-added",
            summary="Администратор и специалист могут скачать Excel-файлы",
            content="В закрытых кабинетах доступны отчеты XLSX по заявкам и выполненным работам.",
            article_type="news",
        ),
    ]
    db.session.add_all(articles)
    db.session.commit()
